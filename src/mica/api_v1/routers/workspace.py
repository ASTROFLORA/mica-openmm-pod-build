"""Workspace session storage router.

VS Code-style workspace sessions persisted to GCS (or local filesystem fallback).
Each user has workspace sessions containing PDB, DCD, PDF, XML assets and
session metadata.

Storage layout (GCS or local):
    workspaces/{session_id}/session.json
    workspaces/{session_id}/assets/{asset_type}/{filename}
"""

from __future__ import annotations

import hashlib
import json
import os
import re
import shutil
import time
import uuid
from datetime import datetime, timezone
from enum import Enum
from pathlib import Path
from typing import Any, Dict, List, Optional

import httpx
from fastapi import APIRouter, Depends, HTTPException, Query, Request
from pydantic import BaseModel, Field

from mica.api_v1.auth import request_identity_dependency
from mica.identity.request_identity import RequestIdentity
from mica.pipelines.knowledge_fabric.document_envelope import DocumentKind, DocumentScanMode
from mica.pipelines.knowledge_fabric.document_scan_service import DocumentScanService

router = APIRouter(prefix="/api/v1/workspace", tags=["workspace"])


def _authenticated_user_id(user: Any) -> str:
    if isinstance(user, RequestIdentity):
        return user.user_id
    return str(user)

# ---------------------------------------------------------------------------
# Asset type enum
# ---------------------------------------------------------------------------

class AssetType(str, Enum):
    pdb = "pdb"
    dcd = "dcd"
    pdf = "pdf"
    xml = "xml"
    document = "document"
    image = "image"
    data = "data"
    other = "other"


ASSET_CONTENT_TYPES: Dict[str, str] = {
    "pdb": "chemical/x-pdb",
    "dcd": "application/octet-stream",
    "pdf": "application/pdf",
    "xml": "application/xml",
    "document": "text/plain; charset=utf-8",
    "image": "application/octet-stream",
    "data": "application/octet-stream",
    "other": "application/octet-stream",
}

# ---------------------------------------------------------------------------
# Request / response models
# ---------------------------------------------------------------------------

class CreateSessionRequest(BaseModel):
    name: str = Field(..., min_length=1, max_length=200, description="Human-readable session name")
    description: Optional[str] = Field(None, max_length=2000)

class AddAssetRequest(BaseModel):
    asset_type: AssetType = Field(..., description="Type of asset")
    name: str = Field(..., min_length=1, max_length=255, description="Display name / filename")
    pdb_id: Optional[str] = Field(None, description="PDB ID for auto-download from RCSB")
    content: Optional[str] = Field(None, description="Inline UTF-8 content for document-like assets")

class SessionMetadata(BaseModel):
    session_id: str
    name: str
    description: Optional[str] = None
    created_at: str
    updated_at: str
    asset_count: int = 0
    assets: List[Dict[str, Any]] = Field(default_factory=list)
    archived: bool = False


class SaveWorkspaceTemplateRequest(BaseModel):
    name: str = Field(..., min_length=1, max_length=200, description="Human-readable template name")
    description: Optional[str] = Field(None, max_length=2000)
    workspace_tools: List[Dict[str, Any]] = Field(default_factory=list)
    active_tool_id: Optional[str] = None


class WorkspaceTemplateRecord(BaseModel):
    template_id: str
    name: str
    description: Optional[str] = None
    created_at: str
    updated_at: str
    workspace_tools: List[Dict[str, Any]] = Field(default_factory=list)
    active_tool_id: Optional[str] = None
    tool_types: List[str] = Field(default_factory=list)


class WorkspaceScanRequest(BaseModel):
    mode: DocumentScanMode = Field(DocumentScanMode.DLM_SECTIONS)
    kb_id: str = Field("", description="Optional KB to associate with this scan")

# ---------------------------------------------------------------------------
# Storage backend abstraction
# ---------------------------------------------------------------------------

_SAFE_NAME_RE = re.compile(r"^[A-Za-z0-9._-]+$")


def _extract_tool_types(workspace_tools: List[Dict[str, Any]]) -> List[str]:
    ordered: List[str] = []
    seen: set[str] = set()
    for tool in workspace_tools:
        if not isinstance(tool, dict):
            continue
        tool_type = tool.get("toolType")
        if not isinstance(tool_type, str):
            continue
        if tool_type in seen:
            continue
        seen.add(tool_type)
        ordered.append(tool_type)
    return ordered


def _sanitize_filename(name: str) -> str:
    """Produce a safe filename from user input."""
    safe = re.sub(r"[^A-Za-z0-9._-]", "_", name.strip())
    if not safe or safe in (".", ".."):
        safe = "unnamed"
    return safe[:255]


def _repo_root() -> Path:
    """Return the repository root (two levels above src/mica/api_v1/routers)."""
    return Path(__file__).resolve().parents[4]


class _LocalBackend:
    """Filesystem-based workspace storage for development."""

    def __init__(self, base_dir: Path) -> None:
        self._base = base_dir

    def _user_dir(self, user_id: str) -> Path:
        digest = hashlib.sha256(user_id.encode()).hexdigest()[:16]
        return self._base / digest

    def _session_dir(self, user_id: str, session_id: str) -> Path:
        return self._user_dir(user_id) / session_id

    def _meta_path(self, user_id: str, session_id: str) -> Path:
        return self._session_dir(user_id, session_id) / "session.json"

    def _template_dir(self, user_id: str) -> Path:
        return self._user_dir(user_id) / "templates"

    def _template_path(self, user_id: str, template_id: str) -> Path:
        return self._template_dir(user_id) / f"{template_id}.json"

    # -- session CRUD -------------------------------------------------------

    def create_session(self, user_id: str, name: str, description: Optional[str]) -> SessionMetadata:
        session_id = uuid.uuid4().hex
        now = datetime.now(timezone.utc).isoformat()
        meta = SessionMetadata(
            session_id=session_id,
            name=name,
            description=description,
            created_at=now,
            updated_at=now,
        )
        sd = self._session_dir(user_id, session_id)
        sd.mkdir(parents=True, exist_ok=True)
        self._write_meta(user_id, session_id, meta)
        return meta

    def list_sessions(self, user_id: str) -> List[SessionMetadata]:
        ud = self._user_dir(user_id)
        if not ud.exists():
            return []
        sessions: List[SessionMetadata] = []
        for child in sorted(ud.iterdir()):
            mp = child / "session.json"
            if mp.is_file():
                try:
                    meta = self._read_meta(mp)
                    if not meta.archived:
                        sessions.append(meta)
                except Exception:
                    continue
        return sessions

    def get_session(self, user_id: str, session_id: str) -> SessionMetadata:
        mp = self._meta_path(user_id, session_id)
        if not mp.is_file():
            raise HTTPException(status_code=404, detail="Session not found")
        return self._read_meta(mp)

    def delete_session(self, user_id: str, session_id: str) -> None:
        meta = self.get_session(user_id, session_id)
        meta.archived = True
        meta.updated_at = datetime.now(timezone.utc).isoformat()
        self._write_meta(user_id, session_id, meta)

    # -- template CRUD ------------------------------------------------------

    def save_template(
        self,
        user_id: str,
        name: str,
        description: Optional[str],
        workspace_tools: List[Dict[str, Any]],
        active_tool_id: Optional[str],
    ) -> WorkspaceTemplateRecord:
        template_dir = self._template_dir(user_id)
        template_dir.mkdir(parents=True, exist_ok=True)
        existing = next(
            (t for t in self.list_templates(user_id) if t.name.strip().lower() == name.strip().lower()),
            None,
        )
        now = datetime.now(timezone.utc).isoformat()
        template = WorkspaceTemplateRecord(
            template_id=existing.template_id if existing else uuid.uuid4().hex,
            name=name.strip(),
            description=description,
            created_at=existing.created_at if existing else now,
            updated_at=now,
            workspace_tools=workspace_tools,
            active_tool_id=active_tool_id,
            tool_types=_extract_tool_types(workspace_tools),
        )
        self._template_path(user_id, template.template_id).write_text(
            template.model_dump_json(indent=2),
            encoding="utf-8",
        )
        return template

    def list_templates(self, user_id: str) -> List[WorkspaceTemplateRecord]:
        template_dir = self._template_dir(user_id)
        if not template_dir.exists():
            return []
        templates: List[WorkspaceTemplateRecord] = []
        for child in sorted(template_dir.glob("*.json")):
            try:
                templates.append(WorkspaceTemplateRecord.model_validate_json(child.read_text(encoding="utf-8")))
            except Exception:
                continue
        templates.sort(key=lambda item: item.updated_at, reverse=True)
        return templates

    def get_template(self, user_id: str, template_id: str) -> WorkspaceTemplateRecord:
        path = self._template_path(user_id, template_id)
        if not path.is_file():
            raise HTTPException(status_code=404, detail="Workspace template not found")
        return WorkspaceTemplateRecord.model_validate_json(path.read_text(encoding="utf-8"))

    def delete_template(self, user_id: str, template_id: str) -> None:
        path = self._template_path(user_id, template_id)
        if not path.is_file():
            raise HTTPException(status_code=404, detail="Workspace template not found")
        path.unlink()

    # -- asset CRUD ---------------------------------------------------------

    def add_asset(
        self, user_id: str, session_id: str, asset_type: str, name: str, data: bytes,
    ) -> Dict[str, Any]:
        meta = self.get_session(user_id, session_id)
        asset_id = uuid.uuid4().hex[:12]
        safe_name = _sanitize_filename(name)
        rel_path = f"assets/{asset_type}/{safe_name}"
        full_path = self._session_dir(user_id, session_id) / rel_path
        full_path.parent.mkdir(parents=True, exist_ok=True)
        full_path.write_bytes(data)

        asset_record: Dict[str, Any] = {
            "asset_id": asset_id,
            "asset_type": asset_type,
            "name": safe_name,
            "path": rel_path,
            "size_bytes": len(data),
            "created_at": datetime.now(timezone.utc).isoformat(),
        }
        meta.assets.append(asset_record)
        meta.asset_count = len(meta.assets)
        meta.updated_at = datetime.now(timezone.utc).isoformat()
        self._write_meta(user_id, session_id, meta)
        return asset_record

    def list_assets(self, user_id: str, session_id: str) -> List[Dict[str, Any]]:
        meta = self.get_session(user_id, session_id)
        return meta.assets

    def get_asset_path(self, user_id: str, session_id: str, asset_id: str) -> Path:
        meta = self.get_session(user_id, session_id)
        for a in meta.assets:
            if a["asset_id"] == asset_id:
                return self._session_dir(user_id, session_id) / a["path"]
        raise HTTPException(status_code=404, detail="Asset not found")

    def get_asset_download_info(
        self, user_id: str, session_id: str, asset_id: str,
    ) -> Dict[str, Any]:
        """Return a local file:// URL (dev) or raise 404."""
        path = self.get_asset_path(user_id, session_id, asset_id)
        if not path.is_file():
            raise HTTPException(status_code=404, detail="Asset file missing from disk")
        return {"url": path.as_uri(), "expires_in": None, "backend": "local"}

    def read_asset_content(
        self, user_id: str, session_id: str, asset_id: str,
    ) -> tuple[bytes, Dict[str, Any]]:
        """Return raw bytes and asset metadata for the given asset."""
        meta = self.get_session(user_id, session_id)
        asset_record: Optional[Dict[str, Any]] = None
        for a in meta.assets:
            if a["asset_id"] == asset_id:
                asset_record = a
                break
        if asset_record is None:
            raise HTTPException(status_code=404, detail="Asset not found")
        path = self._session_dir(user_id, session_id) / asset_record["path"]
        if not path.is_file():
            raise HTTPException(status_code=404, detail="Asset file missing from disk")
        return path.read_bytes(), asset_record

    # -- helpers ------------------------------------------------------------

    def _write_meta(self, user_id: str, session_id: str, meta: SessionMetadata) -> None:
        mp = self._meta_path(user_id, session_id)
        mp.write_text(meta.model_dump_json(indent=2), encoding="utf-8")

    @staticmethod
    def _read_meta(path: Path) -> SessionMetadata:
        return SessionMetadata.model_validate_json(path.read_text(encoding="utf-8"))


class _GCSBackend:
    """GCS-based workspace storage using GCSUserStorage."""

    def __init__(self) -> None:
        from mica.storage.gcs_user_storage import get_storage_manager
        self._storage = get_storage_manager()

    def _blob_prefix(self, session_id: str) -> str:
        return f"workspaces/{session_id}"

    def _meta_blob(self, session_id: str) -> str:
        return f"workspaces/{session_id}/session.json"

    def _asset_blob(self, session_id: str, asset_type: str, filename: str) -> str:
        return f"workspaces/{session_id}/assets/{asset_type}/{filename}"

    def _template_blob(self, template_id: str) -> str:
        return f"workspace_templates/{template_id}.json"

    # -- session CRUD -------------------------------------------------------

    def create_session(self, user_id: str, name: str, description: Optional[str]) -> SessionMetadata:
        session_id = uuid.uuid4().hex
        now = datetime.now(timezone.utc).isoformat()
        meta = SessionMetadata(
            session_id=session_id,
            name=name,
            description=description,
            created_at=now,
            updated_at=now,
        )
        self._write_meta_gcs(user_id, session_id, meta)
        return meta

    def list_sessions(self, user_id: str) -> List[SessionMetadata]:
        bucket_info = self._storage.ensure_bucket(user_id)
        bucket = self._storage.client.bucket(bucket_info.bucket_name)
        blobs = bucket.list_blobs(prefix="workspaces/", delimiter="/")
        sessions: List[SessionMetadata] = []
        # list_blobs with delimiter gives prefixes for "directories"
        prefixes: List[str] = []
        # Must consume iterator to populate prefixes
        for _ in blobs:
            pass
        prefixes = list(blobs.prefixes)
        for pfx in prefixes:
            # pfx looks like "workspaces/<session_id>/"
            sid = pfx.strip("/").split("/")[-1]
            meta_blob = bucket.blob(self._meta_blob(sid))
            if meta_blob.exists():
                try:
                    raw = meta_blob.download_as_text()
                    meta = SessionMetadata.model_validate_json(raw)
                    if not meta.archived:
                        sessions.append(meta)
                except Exception:
                    continue
        return sessions

    def get_session(self, user_id: str, session_id: str) -> SessionMetadata:
        bucket_info = self._storage.ensure_bucket(user_id)
        bucket = self._storage.client.bucket(bucket_info.bucket_name)
        meta_blob = bucket.blob(self._meta_blob(session_id))
        if not meta_blob.exists():
            raise HTTPException(status_code=404, detail="Session not found")
        raw = meta_blob.download_as_text()
        return SessionMetadata.model_validate_json(raw)

    def delete_session(self, user_id: str, session_id: str) -> None:
        meta = self.get_session(user_id, session_id)
        meta.archived = True
        meta.updated_at = datetime.now(timezone.utc).isoformat()
        self._write_meta_gcs(user_id, session_id, meta)

    # -- template CRUD ------------------------------------------------------

    def save_template(
        self,
        user_id: str,
        name: str,
        description: Optional[str],
        workspace_tools: List[Dict[str, Any]],
        active_tool_id: Optional[str],
    ) -> WorkspaceTemplateRecord:
        existing = next(
            (t for t in self.list_templates(user_id) if t.name.strip().lower() == name.strip().lower()),
            None,
        )
        now = datetime.now(timezone.utc).isoformat()
        template = WorkspaceTemplateRecord(
            template_id=existing.template_id if existing else uuid.uuid4().hex,
            name=name.strip(),
            description=description,
            created_at=existing.created_at if existing else now,
            updated_at=now,
            workspace_tools=workspace_tools,
            active_tool_id=active_tool_id,
            tool_types=_extract_tool_types(workspace_tools),
        )

        bucket_info = self._storage.ensure_bucket(user_id)
        bucket = self._storage.client.bucket(bucket_info.bucket_name)
        blob = bucket.blob(self._template_blob(template.template_id))
        blob.upload_from_string(
            template.model_dump_json(indent=2),
            content_type="application/json",
        )
        return template

    def list_templates(self, user_id: str) -> List[WorkspaceTemplateRecord]:
        bucket_info = self._storage.ensure_bucket(user_id)
        bucket = self._storage.client.bucket(bucket_info.bucket_name)
        templates: List[WorkspaceTemplateRecord] = []
        for blob in bucket.list_blobs(prefix="workspace_templates/"):
            if not blob.name.endswith(".json"):
                continue
            try:
                templates.append(WorkspaceTemplateRecord.model_validate_json(blob.download_as_text()))
            except Exception:
                continue
        templates.sort(key=lambda item: item.updated_at, reverse=True)
        return templates

    def get_template(self, user_id: str, template_id: str) -> WorkspaceTemplateRecord:
        bucket_info = self._storage.ensure_bucket(user_id)
        bucket = self._storage.client.bucket(bucket_info.bucket_name)
        blob = bucket.blob(self._template_blob(template_id))
        if not blob.exists():
            raise HTTPException(status_code=404, detail="Workspace template not found")
        return WorkspaceTemplateRecord.model_validate_json(blob.download_as_text())

    def delete_template(self, user_id: str, template_id: str) -> None:
        bucket_info = self._storage.ensure_bucket(user_id)
        bucket = self._storage.client.bucket(bucket_info.bucket_name)
        blob = bucket.blob(self._template_blob(template_id))
        if not blob.exists():
            raise HTTPException(status_code=404, detail="Workspace template not found")
        blob.delete()

    # -- asset CRUD ---------------------------------------------------------

    def add_asset(
        self, user_id: str, session_id: str, asset_type: str, name: str, data: bytes,
    ) -> Dict[str, Any]:
        meta = self.get_session(user_id, session_id)
        asset_id = uuid.uuid4().hex[:12]
        safe_name = _sanitize_filename(name)
        blob_path = self._asset_blob(session_id, asset_type, safe_name)

        bucket_info = self._storage.ensure_bucket(user_id)
        bucket = self._storage.client.bucket(bucket_info.bucket_name)
        blob = bucket.blob(blob_path)
        ct = ASSET_CONTENT_TYPES.get(asset_type, "application/octet-stream")
        blob.upload_from_string(data, content_type=ct)

        asset_record: Dict[str, Any] = {
            "asset_id": asset_id,
            "asset_type": asset_type,
            "name": safe_name,
            "path": blob_path,
            "size_bytes": len(data),
            "created_at": datetime.now(timezone.utc).isoformat(),
        }
        meta.assets.append(asset_record)
        meta.asset_count = len(meta.assets)
        meta.updated_at = datetime.now(timezone.utc).isoformat()
        self._write_meta_gcs(user_id, session_id, meta)
        return asset_record

    def list_assets(self, user_id: str, session_id: str) -> List[Dict[str, Any]]:
        meta = self.get_session(user_id, session_id)
        return meta.assets

    def get_asset_download_info(
        self, user_id: str, session_id: str, asset_id: str,
    ) -> Dict[str, Any]:
        meta = self.get_session(user_id, session_id)
        asset = None
        for a in meta.assets:
            if a["asset_id"] == asset_id:
                asset = a
                break
        if asset is None:
            raise HTTPException(status_code=404, detail="Asset not found")
        url = self._storage.signed_url(
            user_id=user_id,
            object_path=asset["path"],
            method="GET",
            expires_seconds=900,
        )
        return {"url": url, "expires_in": 900, "backend": "gcs"}

    def read_asset_content(
        self, user_id: str, session_id: str, asset_id: str,
    ) -> tuple[bytes, Dict[str, Any]]:
        """Download raw bytes from GCS and return with asset metadata."""
        meta = self.get_session(user_id, session_id)
        asset_record: Optional[Dict[str, Any]] = None
        for a in meta.assets:
            if a["asset_id"] == asset_id:
                asset_record = a
                break
        if asset_record is None:
            raise HTTPException(status_code=404, detail="Asset not found")

        bucket_info = self._storage.ensure_bucket(user_id)
        bucket = self._storage.client.bucket(bucket_info.bucket_name)
        blob = bucket.blob(asset_record["path"])
        if not blob.exists():
            raise HTTPException(status_code=404, detail="Asset file missing from GCS")
        data = blob.download_as_bytes()
        return data, asset_record

    # -- helpers ------------------------------------------------------------

    def _write_meta_gcs(self, user_id: str, session_id: str, meta: SessionMetadata) -> None:
        bucket_info = self._storage.ensure_bucket(user_id)
        bucket = self._storage.client.bucket(bucket_info.bucket_name)
        blob = bucket.blob(self._meta_blob(session_id))
        blob.upload_from_string(meta.model_dump_json(indent=2), content_type="application/json")


# ---------------------------------------------------------------------------
# Backend selection (lazy singleton)
# ---------------------------------------------------------------------------

_backend: Optional[_LocalBackend | _GCSBackend] = None


def _workspace_backend_mode() -> str:
    mode = (os.getenv("MICA_WORKSPACE_BACKEND") or "auto").strip().lower()
    if mode in {"gcs", "local", "auto"}:
        return mode
    return "auto"


def _explicit_google_credentials_available() -> bool:
    creds_path = (os.getenv("GOOGLE_APPLICATION_CREDENTIALS") or os.getenv("MICA_GOOGLE_APPLICATION_CREDENTIALS") or "").strip()
    if creds_path:
        candidate = Path(creds_path).expanduser()
        if candidate.exists() and candidate.is_file():
            return True

    inline_json = (os.getenv("GOOGLE_APPLICATION_CREDENTIALS_JSON") or os.getenv("MICA_GOOGLE_APPLICATION_CREDENTIALS_JSON") or "").strip()
    return bool(inline_json)


def _local_backend() -> _LocalBackend:
    global _backend
    local_root = _repo_root() / ".workspace_sessions"
    local_root.mkdir(parents=True, exist_ok=True)
    _backend = _LocalBackend(local_root)
    return _backend


def _get_document_scan_service(request: Request) -> DocumentScanService:
    svc = getattr(request.app.state, "document_scan_service", None)
    if svc is None:
        svc = DocumentScanService(store=getattr(request.app.state, "kb_store", None))
        request.app.state.document_scan_service = svc
    return svc


def _get_backend() -> _LocalBackend | _GCSBackend:
    global _backend
    if _backend is not None:
        return _backend

    # Ensure dotenv is loaded so GCP_PROJECT / GOOGLE_APPLICATION_CREDENTIALS
    # are available even when called before the main FastAPI startup event.
    try:
        from mica.config.dotenv_loader import seed_env_from_dotenv
        seed_env_from_dotenv()
    except Exception:
        pass

    mode = _workspace_backend_mode()
    has_project = bool(os.getenv("GCP_PROJECT") or os.getenv("GOOGLE_CLOUD_PROJECT"))
    has_creds = _explicit_google_credentials_available()

    # Avoid triggering google credential discovery on the request path when no
    # explicit credential material is present. Railway/other non-GCP runtimes can
    # otherwise block for tens of seconds before falling back to local storage.
    if mode == "gcs" and has_project and has_creds:
        try:
            _backend = _GCSBackend()
            return _backend
        except Exception:
            return _local_backend()

    if mode == "auto" and has_project and has_creds:
        try:
            _backend = _GCSBackend()
            return _backend
        except Exception:
            return _local_backend()

    return _local_backend()


def workspace_backend_status() -> dict:
    """Return workspace backend type for health checks."""
    try:
        if isinstance(_backend, _GCSBackend):
            return {"backend": "gcs", "status": "ok"}
        if isinstance(_backend, _LocalBackend):
            return {"backend": "local", "status": "ok"}

        mode = _workspace_backend_mode()
        has_project = bool(os.getenv("GCP_PROJECT") or os.getenv("GOOGLE_CLOUD_PROJECT"))
        has_creds = _explicit_google_credentials_available()
        if mode == "gcs":
            if has_project and has_creds:
                return {"backend": "gcs", "status": "configured"}
            return {"backend": "local", "status": "degraded", "reason": "gcs_credentials_missing"}
        if mode == "auto" and has_project and has_creds:
            return {"backend": "gcs", "status": "configured"}
        if has_project and not has_creds:
            return {"backend": "local", "status": "ok", "reason": "gcs_credentials_missing"}
        return {"backend": "local", "status": "ok"}
    except Exception as exc:
        return {"backend": "unknown", "status": "error", "error": str(exc)}


# ---------------------------------------------------------------------------
# PDB auto-download helper
# ---------------------------------------------------------------------------

_RCSB_URL = "https://files.rcsb.org/download/{pdb_id}.pdb"
_PDB_ID_RE = re.compile(r"^[A-Za-z0-9]{4}$")


def _fetch_pdb(pdb_id: str) -> bytes:
    pdb_id = pdb_id.strip().upper()
    if not _PDB_ID_RE.match(pdb_id):
        raise HTTPException(status_code=400, detail=f"Invalid PDB ID: {pdb_id}")
    url = _RCSB_URL.format(pdb_id=pdb_id)
    try:
        resp = httpx.get(url, timeout=30.0, follow_redirects=True)
        resp.raise_for_status()
        return resp.content
    except httpx.HTTPStatusError as exc:
        raise HTTPException(
            status_code=502,
            detail=f"RCSB download failed for {pdb_id}: HTTP {exc.response.status_code}",
        )
    except httpx.RequestError as exc:
        raise HTTPException(status_code=502, detail=f"RCSB download error: {exc}")


# ---------------------------------------------------------------------------
# Routes
# ---------------------------------------------------------------------------

@router.post("/sessions", status_code=201)
def create_session(
    payload: CreateSessionRequest,
    user: Any = Depends(request_identity_dependency),
):
    """Create a new workspace session."""
    backend = _get_backend()
    user_id = _authenticated_user_id(user)
    meta = backend.create_session(user_id, payload.name, payload.description)
    return {"ok": True, "session": meta.model_dump()}


@router.get("/sessions")
def list_sessions(user: Any = Depends(request_identity_dependency)):
    """List all active workspace sessions for the user."""
    backend = _get_backend()
    user_id = _authenticated_user_id(user)
    sessions = backend.list_sessions(user_id)
    return {"ok": True, "sessions": [s.model_dump() for s in sessions]}


@router.get("/sessions/{session_id}")
def get_session(session_id: str, user: Any = Depends(request_identity_dependency)):
    """Get workspace session details including assets."""
    backend = _get_backend()
    user_id = _authenticated_user_id(user)
    meta = backend.get_session(user_id, session_id)
    return {"ok": True, "session": meta.model_dump()}


@router.post("/templates", status_code=201)
def save_workspace_template(
    payload: SaveWorkspaceTemplateRequest,
    user: Any = Depends(request_identity_dependency),
):
    """Create or update a named workspace template for the authenticated user."""
    backend = _get_backend()
    user_id = _authenticated_user_id(user)
    template = backend.save_template(
        user_id=user_id,
        name=payload.name,
        description=payload.description,
        workspace_tools=payload.workspace_tools,
        active_tool_id=payload.active_tool_id,
    )
    return {"ok": True, "template": template.model_dump()}


@router.get("/templates")
def list_workspace_templates(user: Any = Depends(request_identity_dependency)):
    """List saved workspace templates for the authenticated user."""
    backend = _get_backend()
    user_id = _authenticated_user_id(user)
    templates = backend.list_templates(user_id)
    return {"ok": True, "templates": [template.model_dump() for template in templates]}


@router.get("/templates/{template_id}")
def get_workspace_template(template_id: str, user: Any = Depends(request_identity_dependency)):
    """Get the serialized workspace state for a saved template."""
    backend = _get_backend()
    user_id = _authenticated_user_id(user)
    template = backend.get_template(user_id, template_id)
    return {"ok": True, "template": template.model_dump()}


@router.delete("/templates/{template_id}")
def delete_workspace_template(template_id: str, user: Any = Depends(request_identity_dependency)):
    """Delete a saved workspace template for the authenticated user."""
    backend = _get_backend()
    user_id = _authenticated_user_id(user)
    backend.delete_template(user_id, template_id)
    return {"ok": True, "detail": "Workspace template deleted"}


@router.post("/sessions/{session_id}/assets", status_code=201)
def add_asset(
    session_id: str,
    payload: AddAssetRequest,
    user: Any = Depends(request_identity_dependency),
):
    """Add an asset to a workspace session.

    If ``asset_type`` is ``pdb`` and ``pdb_id`` is provided, the PDB file is
    auto-downloaded from RCSB. If ``content`` is provided, it is stored as a
    UTF-8 text asset. Otherwise a zero-byte placeholder is created.
    """
    _MAX_ASSET_BYTES = 50 * 1024 * 1024  # 50 MB

    backend = _get_backend()
    user_id = _authenticated_user_id(user)
    # Ensure session exists
    backend.get_session(user_id, session_id)

    name = payload.name
    data: bytes = b""

    if payload.pdb_id:
        data = _fetch_pdb(payload.pdb_id)
        if not name.lower().endswith(".pdb"):
            name = f"{payload.pdb_id.upper()}.pdb"
    elif payload.content is not None:
        data = payload.content.encode("utf-8")

    if len(data) > _MAX_ASSET_BYTES:
        raise HTTPException(status_code=413, detail="Asset exceeds 50 MB limit")

    asset = backend.add_asset(
        user_id=user_id,
        session_id=session_id,
        asset_type=payload.asset_type.value,
        name=name,
        data=data,
    )
    return {"ok": True, "asset": asset}


@router.get("/sessions/{session_id}/assets")
def list_assets(session_id: str, user: Any = Depends(request_identity_dependency)):
    """List all assets in a workspace session."""
    backend = _get_backend()
    user_id = _authenticated_user_id(user)
    assets = backend.list_assets(user_id, session_id)
    return {"ok": True, "assets": assets}


@router.get("/sessions/{session_id}/assets/{asset_id}/download")
def download_asset(
    session_id: str,
    asset_id: str,
    user: Any = Depends(request_identity_dependency),
):
    """Get a download URL (signed GCS URL or local file URI) for an asset."""
    backend = _get_backend()
    user_id = _authenticated_user_id(user)
    info = backend.get_asset_download_info(user_id, session_id, asset_id)
    return {"ok": True, **info}


@router.delete("/sessions/{session_id}")
def delete_session(session_id: str, user: Any = Depends(request_identity_dependency)):
    """Archive (soft-delete) a workspace session."""
    backend = _get_backend()
    user_id = _authenticated_user_id(user)
    backend.delete_session(user_id, session_id)
    return {"ok": True, "detail": "Session archived"}


# ---------------------------------------------------------------------------
# Document text extraction helpers
# ---------------------------------------------------------------------------

_MAX_PDF_PAGES = 60
_MAX_EXTRACT_CHARS = 80_000  # Prevent overwhelming the LLM context


def _extract_text_from_pdf_bytes(data: bytes, *, max_pages: int = _MAX_PDF_PAGES) -> Optional[str]:
    """Best-effort PDF text extraction from raw bytes (no temp file needed for fitz)."""
    # Prefer PyMuPDF (fitz)
    try:
        import fitz  # type: ignore

        doc = fitz.open(stream=data, filetype="pdf")
        parts: List[str] = []
        for i, page in enumerate(doc):
            if i >= max_pages:
                break
            txt = page.get_text("text") or ""
            if txt:
                parts.append(txt)
        doc.close()
        text = "\n".join(parts).strip()
        return text or None
    except Exception:
        pass

    # Fallback: pypdf (needs BytesIO)
    try:
        import io
        from pypdf import PdfReader  # type: ignore

        reader = PdfReader(io.BytesIO(data))
        parts2: List[str] = []
        for i, page in enumerate(reader.pages):
            if i >= max_pages:
                break
            try:
                parts2.append(page.extract_text() or "")
            except Exception:
                parts2.append("")
        text = "\n".join(parts2).strip()
        return text or None
    except Exception:
        return None


def _extract_text_from_docx_bytes(data: bytes) -> Optional[str]:
    """Extract text from a DOCX file bytes using python-docx."""
    try:
        import io
        from docx import Document  # type: ignore

        doc = Document(io.BytesIO(data))
        parts: List[str] = []
        for para in doc.paragraphs:
            if para.text:
                parts.append(para.text)
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = "\t".join(cell.text for cell in row.cells if cell.text)
                if row_text:
                    parts.append(row_text)
        text = "\n".join(parts).strip()
        return text or None
    except Exception:
        return None


def _extract_text_from_bytes(data: bytes, filename: str, asset_type: str) -> Dict[str, Any]:
    """Dispatch text extraction based on file type. Returns dict with text + format."""
    lower_name = filename.lower()

    # PDF
    if lower_name.endswith(".pdf") or asset_type == "pdf":
        text = _extract_text_from_pdf_bytes(data)
        if text:
            if len(text) > _MAX_EXTRACT_CHARS:
                return {"text": text[:_MAX_EXTRACT_CHARS], "format": "pdf", "truncated": True, "total_chars": len(text)}
            return {"text": text, "format": "pdf", "truncated": False}
        return {"text": None, "format": "pdf", "error": "Could not extract text from PDF (scanned image or empty)"}

    # DOCX
    if lower_name.endswith(".docx"):
        text = _extract_text_from_docx_bytes(data)
        if text:
            if len(text) > _MAX_EXTRACT_CHARS:
                return {"text": text[:_MAX_EXTRACT_CHARS], "format": "docx", "truncated": True, "total_chars": len(text)}
            return {"text": text, "format": "docx", "truncated": False}
        return {"text": None, "format": "docx", "error": "Could not extract text from DOCX"}

    # Plain text formats
    if lower_name.endswith((".txt", ".csv", ".xml", ".json", ".md", ".log", ".pdb", ".cif")):
        try:
            text = data.decode("utf-8", errors="replace")
            fmt = lower_name.rsplit(".", 1)[-1] if "." in lower_name else "text"
            if len(text) > _MAX_EXTRACT_CHARS:
                return {"text": text[:_MAX_EXTRACT_CHARS], "format": fmt, "truncated": True, "total_chars": len(text)}
            return {"text": text, "format": fmt, "truncated": False}
        except Exception:
            return {"text": None, "format": "text", "error": "Decoding failed"}

    return {"text": None, "format": "binary", "error": f"Unsupported format for text extraction: {lower_name}"}


# ---------------------------------------------------------------------------
# Read document route
# ---------------------------------------------------------------------------

@router.get("/sessions/{session_id}/assets/{asset_id}/read")
def read_workspace_document(
    session_id: str,
    asset_id: str,
    max_pages: int = Query(default=_MAX_PDF_PAGES, ge=1, le=200),
    user: Any = Depends(request_identity_dependency),
):
    """Read and extract text from a workspace document (PDF, DOCX, plain text).

    Returns the extracted text content rather than a download URL.
    """
    backend = _get_backend()
    user_id = _authenticated_user_id(user)
    data, asset_record = backend.read_asset_content(user_id, session_id, asset_id)
    result = _extract_text_from_bytes(data, asset_record["name"], asset_record.get("asset_type", "other"))
    result["asset"] = {
        "asset_id": asset_record["asset_id"],
        "name": asset_record["name"],
        "asset_type": asset_record.get("asset_type", "other"),
        "size_bytes": asset_record.get("size_bytes", len(data)),
    }
    result["ok"] = True
    return result


@router.post("/sessions/{session_id}/assets/{asset_id}/scan")
async def scan_workspace_document(
    session_id: str,
    asset_id: str,
    payload: WorkspaceScanRequest,
    request: Request,
    user: Any = Depends(request_identity_dependency),
):
    """Scan a workspace asset into sections/entities/candidate claims."""
    backend = _get_backend()
    user_id = _authenticated_user_id(user)
    data, asset_record = backend.read_asset_content(user_id, session_id, asset_id)
    extracted = _extract_text_from_bytes(data, asset_record["name"], asset_record.get("asset_type", "other"))
    text = str(extracted.get("text") or "").strip()
    scan_service = _get_document_scan_service(request)
    result = await scan_service.create_scan(
        title=str(asset_record.get("name") or asset_id),
        text=text,
        mode=payload.mode,
        document_kind=DocumentKind.WORKSPACE_ASSET,
        owner_id=user_id,
        workspace_id=session_id,
        asset_id=asset_id,
        kb_id=payload.kb_id,
        provider=str(extracted.get("format") or asset_record.get("asset_type") or "workspace"),
        acquisition_type="workspace_asset",
        metadata={
            "asset_name": asset_record.get("name"),
            "asset_type": asset_record.get("asset_type"),
            "size_bytes": asset_record.get("size_bytes", len(data)),
            "extract_format": extracted.get("format"),
            "truncated": extracted.get("truncated", False),
        },
    )
    return {"ok": True, **result.model_dump()}


@router.get("/scans/{scan_id}")
async def get_workspace_scan_status(
    scan_id: str,
    request: Request,
    _user: Any = Depends(request_identity_dependency),
):
    """Return the current status/result for a workspace document scan."""
    scan_service = _get_document_scan_service(request)
    try:
        result = await scan_service.get_scan(scan_id)
    except KeyError:
        raise HTTPException(status_code=404, detail="Scan not found")
    return {"ok": True, **result.model_dump()}
